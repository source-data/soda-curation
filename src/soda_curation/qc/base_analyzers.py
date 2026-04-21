"""Base QC test analyzers."""

import json
import logging
from abc import ABC, abstractmethod
from copy import deepcopy
from typing import Any, Dict, List, Optional, Tuple, Union

from pydantic import BaseModel

from ..pipeline.ai_observability import safe_excerpt, summarize_text
from .model_api import ModelAPI
from .prompt_registry import registry

logger = logging.getLogger(__name__)
SUPPORTED_AI_PROVIDERS = {"openai", "anthropic", "gemini"}


class BaseQCAnalyzer(ABC):
    """Base class for all QC analyzers."""

    def __init__(self, config: Dict):
        """Initialize the analyzer."""
        self.config = config
        self.model_api = ModelAPI(config)
        self.test_name = self.__class__.__name__.replace("Analyzer", "").lower()

        # Get metadata from registry for traceability
        self.metadata = registry.get_prompt_metadata(self.test_name)

        # Get the dynamically generated model
        self.result_model = registry.get_pydantic_model(self.test_name)

    def get_test_config(self) -> Dict:
        """Get the test-specific configuration."""
        # Look in default.pipeline if available
        if "default" in self.config and "pipeline" in self.config["default"]:
            pipeline_config = self.config["default"]["pipeline"]
            if self.test_name in pipeline_config:
                return pipeline_config[self.test_name]

        # Fallback to root level
        return self.config.get(self.test_name, {})

    def _get_provider_config(self, test_config: Dict) -> Dict:
        """Return the provider-specific config section for the active AI provider."""
        provider = self.config.get("ai_provider", "openai").lower()
        if provider not in SUPPORTED_AI_PROVIDERS:
            raise ValueError(
                f"Unsupported ai_provider '{provider}' for QC test '{self.test_name}'. "
                f"Expected one of {sorted(SUPPORTED_AI_PROVIDERS)}."
            )

        provider_keys = ("openai", "anthropic", "gemini")
        configured_provider_blocks = [
            key for key in provider_keys if key in test_config
        ]
        assert len(configured_provider_blocks) <= 1, (
            f"QC configuration error for test '{self.test_name}': provider blocks "
            f"{configured_provider_blocks} are defined together. Define only one provider "
            "per test."
        )

        configured_provider = (
            configured_provider_blocks[0] if configured_provider_blocks else None
        )
        if configured_provider and configured_provider != provider:
            raise ValueError(
                f"QC provider mismatch for test '{self.test_name}': "
                f"configured_provider='{configured_provider}', ai_provider='{provider}'."
            )

        if configured_provider == provider:
            provider_config = deepcopy(test_config[provider])
        elif "default" in self.config and provider in self.config["default"]:
            provider_config = deepcopy(self.config["default"][provider])
        else:
            # Safe baseline defaults when no provider block exists.
            if provider == "anthropic":
                provider_config = {
                    "model": "claude-sonnet-4-6",
                    "temperature": 0.1,
                    "max_tokens": 4096,
                }
            elif provider == "gemini":
                provider_config = {
                    "model": "gemini-2.5-flash",
                    "temperature": 0.1,
                    "max_tokens": 2048,
                }
            else:
                provider_config = {
                    "model": "gpt-4o",
                    "temperature": 0.1,
                    "json_mode": True,
                }

        provider_config.setdefault("prompts", {})
        return provider_config

    def _apply_langfuse_runtime_hints(self, provider_config: Dict[str, Any]) -> None:
        """Merge optional prompt-level runtime hints from Langfuse config."""
        runtime_hints = registry.get_runtime_config(self.test_name)
        if not runtime_hints:
            return
        _deep_merge_dict(provider_config, runtime_hints)

    @abstractmethod
    def analyze(self, *args, **kwargs) -> Tuple[bool, Any]:
        """Run the analysis and return results."""
        pass

    def create_empty_result(self) -> Dict:
        """Create an empty result object."""
        return {"outputs": []}


class PanelQCAnalyzer(BaseQCAnalyzer):
    """Base class for panel-level QC tests."""

    def analyze_figure(
        self,
        figure_label: str,
        encoded_image: str,
        figure_caption: str,
        expected_panels: list = None,
    ) -> Tuple[bool, Any]:
        """Analyze a figure and return results."""
        try:
            if not encoded_image:
                logger.warning(
                    "Skipping panel QC call because encoded image is empty",
                    extra={
                        "operation": "qc.panel_check",
                        "test_name": self.test_name,
                        "figure_label": figure_label,
                        "severity": "recoverable",
                        "reason": "empty_image_payload",
                    },
                )
                return False, self.create_empty_result()

            # Get API config from main config or use defaults
            test_config = self.get_test_config()
            provider_config = self._get_provider_config(test_config)

            # Set system prompt with one from registry
            provider_config["prompts"]["system"] = registry.get_prompt(self.test_name)
            provider_config["prompts"]["user"] = (
                registry.get_user_prompt(self.test_name)
                or provider_config["prompts"].get("user")
                or "$figure_caption"
            )
            self._apply_langfuse_runtime_hints(provider_config)

            # Set user prompt to include the figure caption if not already set
            if (
                "user" not in provider_config["prompts"]
                or not provider_config["prompts"]["user"]
            ):
                provider_config["prompts"]["user"] = "Figure caption:\n$figure_caption"

            # Call the model with the correct parameters
            response = self.model_api.generate_response(
                encoded_image=encoded_image,
                caption=figure_caption,
                prompt_config=provider_config,
                response_type=self.result_model,
                expected_panels=expected_panels,
                operation="qc.panel_check",
                context={
                    "test_name": self.test_name,
                    "figure_label": figure_label,
                    "expected_panel_count": len(expected_panels or []),
                    "figure_caption_summary": summarize_text(figure_caption),
                },
            )

            # Process and validate the response
            result = self.process_response(response)

            # Filter out panels with labels not in expected_panels
            result = self._filter_valid_panels(result, expected_panels)

            # Check if the test passed
            passed = self.check_test_passed(result)

            return passed, result

        except ValueError as e:
            logger.error(
                "Panel QC call failed due to invalid request setup",
                extra={
                    "operation": "qc.panel_check",
                    "test_name": self.test_name,
                    "figure_label": figure_label,
                    "severity": "critical",
                    "reason": "invalid_request_setup",
                    "error": str(e),
                },
            )
            return False, self.create_empty_result()
        except Exception as e:
            logger.error(
                "Panel QC call failed with unexpected error",
                extra={
                    "operation": "qc.panel_check",
                    "test_name": self.test_name,
                    "figure_label": figure_label,
                    "severity": "recoverable",
                    "reason": "unexpected_analyzer_error",
                    "error": str(e),
                },
            )
            return False, self.create_empty_result()

    @staticmethod
    def _filter_valid_panels(
        result: Dict, expected_panels: Optional[List[str]]
    ) -> Dict:
        """Filter response outputs to only include panels with valid labels.

        Args:
            result: Parsed response dict with an "outputs" list.
            expected_panels: Allowed panel labels for this figure.

        Returns:
            The result dict with invalid-label outputs removed.
        """
        if not expected_panels or not isinstance(result, dict):
            return result

        outputs = result.get("outputs", [])
        if not outputs:
            return result

        valid_outputs = []
        for panel in outputs:
            # Extract panel_label from dict or object
            if isinstance(panel, dict):
                label = panel.get("panel_label")
            elif hasattr(panel, "panel_label"):
                label = panel.panel_label
            else:
                label = None

            if label in expected_panels:
                valid_outputs.append(panel)
            else:
                logger.warning(
                    f"Discarded panel with invalid label '{label}' "
                    f"(expected one of {expected_panels})"
                )

        result["outputs"] = valid_outputs
        return result

    def process_response(self, response: Any) -> Any:
        """Process the response from the model API."""
        # If response is a string, try to parse it
        if isinstance(response, str):
            try:
                response_data = json.loads(response)
            except json.JSONDecodeError:
                logger.error(
                    "Failed to parse panel QC response as JSON",
                    extra={
                        "operation": "qc.panel_check",
                        "test_name": self.test_name,
                        "response_excerpt": safe_excerpt(response),
                    },
                )
                return self.create_empty_result()
        else:
            response_data = response

        # Handle both direct outputs and nested outputs
        if isinstance(response_data, dict):
            if "outputs" in response_data:
                return response_data
            else:
                # Wrap in outputs structure if not already present
                return {"outputs": [response_data]}

        # Return empty result if we can't process the response
        return self.create_empty_result()

    def check_test_passed(self, result: Any) -> bool:
        """Check if the test passed based on the result."""
        # Default implementation
        if not result or not isinstance(result, dict) or "outputs" not in result:
            return False

        # If any panel has issues, the test fails
        for panel in result.get("outputs", []):
            if not self.check_panel_passed(panel):
                return False

        # If we get here, all panels passed
        return True

    def check_panel_passed(self, panel: Any) -> bool:
        """Check if an individual panel passed the test."""
        # Default implementation considers any non-empty panel as passed
        return panel is not None and len(panel) > 0


class FigureQCAnalyzer(BaseQCAnalyzer):
    """Base class for figure-level QC tests."""

    def analyze_figure(
        self, figure_label: str, encoded_image: str, figure_caption: str
    ) -> Tuple[bool, Any]:
        """Analyze a whole figure and return results."""
        try:
            if not encoded_image:
                logger.warning(
                    "Skipping figure QC call because encoded image is empty",
                    extra={
                        "operation": "qc.figure_check",
                        "test_name": self.test_name,
                        "figure_label": figure_label,
                        "severity": "recoverable",
                        "reason": "empty_image_payload",
                    },
                )
                return False, self.create_empty_result()

            # Get API config from main config or use defaults
            test_config = self.get_test_config()
            provider_config = self._get_provider_config(test_config)

            # Set system prompt with one from registry
            provider_config["prompts"]["system"] = registry.get_prompt(self.test_name)
            provider_config["prompts"]["user"] = (
                registry.get_user_prompt(self.test_name)
                or provider_config["prompts"].get("user")
                or "$figure_caption"
            )
            self._apply_langfuse_runtime_hints(provider_config)

            # Set user prompt to include the figure caption if not already set
            if (
                "user" not in provider_config["prompts"]
                or not provider_config["prompts"]["user"]
            ):
                provider_config["prompts"]["user"] = "Figure caption:\n$figure_caption"

            # Call the model with the correct parameters
            response = self.model_api.generate_response(
                encoded_image=encoded_image,
                caption=figure_caption,
                prompt_config=provider_config,
                response_type=self.result_model,
                operation="qc.figure_check",
                context={
                    "test_name": self.test_name,
                    "figure_label": figure_label,
                    "figure_caption_summary": summarize_text(figure_caption),
                },
            )

            # Process and validate the response
            result = self.process_response(response)

            # Check if the test passed
            passed = self.check_test_passed(result)

            return passed, result

        except ValueError as e:
            logger.error(
                "Figure QC call failed due to invalid request setup",
                extra={
                    "operation": "qc.figure_check",
                    "test_name": self.test_name,
                    "figure_label": figure_label,
                    "severity": "critical",
                    "reason": "invalid_request_setup",
                    "error": str(e),
                },
            )
            return False, self.create_empty_result()
        except Exception as e:
            logger.error(
                "Figure QC call failed with unexpected error",
                extra={
                    "operation": "qc.figure_check",
                    "test_name": self.test_name,
                    "figure_label": figure_label,
                    "severity": "recoverable",
                    "reason": "unexpected_analyzer_error",
                    "error": str(e),
                },
            )
            return False, self.create_empty_result()

    def process_response(self, response: Any) -> Any:
        """Process the response from the model API."""
        if not response:
            return self.create_empty_result()

        # If response is a JSON string, parse it
        if isinstance(response, str):
            try:
                import json

                parsed_response = json.loads(response)
                return parsed_response
            except json.JSONDecodeError:
                logger.error(
                    "Failed to parse figure QC response as JSON",
                    extra={
                        "operation": "qc.figure_check",
                        "test_name": self.test_name,
                        "response_excerpt": safe_excerpt(response),
                    },
                )
                return self.create_empty_result()

        return response

    def check_test_passed(self, result: Any) -> bool:
        """Check if the test passed based on the result."""
        # Figure-level implementation
        return bool(result)


class ManuscriptQCAnalyzer(BaseQCAnalyzer):
    """Base class for manuscript-level QC tests."""

    def analyze_manuscript(
        self, zip_structure: Any, word_file_path: str = None
    ) -> Tuple[bool, Any]:
        """Analyze an entire manuscript and return results."""
        try:
            # Get API config from main config or use defaults
            test_config = self.get_test_config()
            provider_config = self._get_provider_config(test_config)

            # Set system prompt with one from registry
            provider_config["prompts"]["system"] = registry.get_prompt(self.test_name)
            provider_config["prompts"]["user"] = (
                registry.get_user_prompt(self.test_name)
                or provider_config["prompts"].get("user")
                or "$manuscript_text"
            )
            self._apply_langfuse_runtime_hints(provider_config)

            # Set user prompt template: manuscript text is always the user message
            if "user" not in provider_config["prompts"]:
                provider_config["prompts"]["user"] = "$manuscript_text"

            # Extract word file content from zip_structure or provided path
            word_file_content = self.extract_word_file_content(
                zip_structure, word_file_path
            )

            # Call the model with the correct parameters
            response = self.model_api.generate_response(
                prompt_config=provider_config,
                response_type=self.result_model,
                word_file_content=word_file_content,
                operation="qc.manuscript_check",
                context={
                    "test_name": self.test_name,
                    "word_file_summary": summarize_text(word_file_content),
                },
            )

            # Process and validate the response
            result = self.process_response(response)

            # Check if the test passed
            passed = self.check_test_passed(result)

            return passed, result

        except ValueError as e:
            logger.error(
                "Manuscript QC call failed due to invalid request setup",
                extra={
                    "operation": "qc.manuscript_check",
                    "test_name": self.test_name,
                    "severity": "critical",
                    "reason": "invalid_request_setup",
                    "error": str(e),
                },
            )
            return False, self.create_empty_result()
        except Exception as e:
            logger.error(
                "Manuscript QC call failed with unexpected error",
                extra={
                    "operation": "qc.manuscript_check",
                    "test_name": self.test_name,
                    "severity": "recoverable",
                    "reason": "unexpected_analyzer_error",
                    "error": str(e),
                },
            )
            return False, self.create_empty_result()

    def extract_word_file_content(
        self, zip_structure: Any, word_file_path: str = None
    ) -> str:
        """Extract text content from word file in the manuscript structure."""
        try:
            # Use pre-extracted manuscript text stored in zip_structure (preferred)
            if (
                not word_file_path
                and zip_structure
                and hasattr(zip_structure, "manuscript_text")
                and zip_structure.manuscript_text
            ):
                return zip_structure.manuscript_text

            # Import python-docx for reading Word documents
            from docx import Document

            # Determine the word file path to use
            docx_path = None

            if word_file_path:
                # Use provided word file path
                docx_path = word_file_path
            elif (
                zip_structure
                and hasattr(zip_structure, "_full_docx")
                and zip_structure._full_docx
            ):
                # Use full path from zip structure
                docx_path = zip_structure._full_docx
            elif (
                zip_structure and hasattr(zip_structure, "docx") and zip_structure.docx
            ):
                # Fallback to relative path (might need to be joined with extract directory)
                docx_path = zip_structure.docx

            if not docx_path:
                logger.warning("No word file path found in zip structure")
                return "No word file available for analysis"

            # Check if file exists
            import os

            if not os.path.exists(docx_path):
                logger.warning(f"Word file not found at path: {docx_path}")
                return f"Word file not found at path: {docx_path}"

            # Extract text from the Word document
            doc = Document(docx_path)

            # Extract all text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only include non-empty paragraphs
                    text_content.append(paragraph.text.strip())

            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())

            # Join all text with newlines
            full_text = "\n".join(text_content)

            if not full_text.strip():
                logger.warning("No text content found in word file")
                return "No text content found in word file"

            logger.info(
                f"Successfully extracted {len(full_text)} characters from word file"
            )
            return full_text

        except ImportError:
            logger.error("python-docx library not available for reading Word documents")
            return "Error: python-docx library not available"
        except Exception as e:
            logger.error(f"Error extracting word file content: {str(e)}")
            return f"Error extracting word file content: {str(e)}"

    def extract_manuscript_text(self, zip_structure: Any) -> str:
        """Extract relevant text from the manuscript structure. (Legacy method - kept for compatibility)"""
        # Delegate to the new word file extraction method
        return self.extract_word_file_content(zip_structure)

    def process_response(self, response: Any) -> Any:
        """Process the response from the model API."""
        if not response:
            return self.create_empty_result()

        # If response is a JSON string, parse it
        if isinstance(response, str):
            try:
                import json

                parsed_response = json.loads(response)
                return parsed_response
            except json.JSONDecodeError:
                logger.error(
                    "Failed to parse manuscript QC response as JSON",
                    extra={
                        "operation": "qc.manuscript_check",
                        "test_name": self.test_name,
                        "response_excerpt": safe_excerpt(response),
                    },
                )
                return self.create_empty_result()

        return response

    def check_test_passed(self, result: Any) -> bool:
        """Check if the test passed based on the result."""
        # Manuscript-level implementation
        return bool(result)


def _deep_merge_dict(base: Dict[str, Any], extra: Dict[str, Any]) -> None:
    """In-place deep merge used for runtime-hint overlays."""
    for key, value in extra.items():
        if key in base and isinstance(base[key], dict) and isinstance(value, dict):
            _deep_merge_dict(base[key], value)
            continue
        base[key] = deepcopy(value)
