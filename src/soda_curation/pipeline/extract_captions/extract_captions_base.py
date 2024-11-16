"""
This module provides the base class for figure caption extraction.

It defines the abstract base class that all specific caption extractor implementations
should inherit from, ensuring a consistent interface across different extraction methods.
"""

import json
import logging
import re
import unicodedata
from abc import ABC, abstractmethod
from typing import Dict, Optional, Tuple

from docx import Document
from rouge_score import rouge_scorer

from ..manuscript_structure.manuscript_structure import ZipStructure

logger = logging.getLogger(__name__)

def normalize_text(text: str) -> str:
    """
    Normalize text for comparison by converting to ASCII and handling scientific notation.
    
    Args:
        text (str): Input text to normalize
        
    Returns:
        str: Normalized text ready for comparison
    """
    # First decompose Unicode characters
    text = unicodedata.normalize('NFKD', text)
    
    # Remove combining characters (accents, diacritics)
    text = ''.join(c for c in text if not unicodedata.combining(c))
    
    # Common scientific symbol replacements
    replacements = {
        '±': 'plus/minus',
        '→': '->',
        '←': '<-',
        '°': ' degrees ',
        'µ': 'micro',
        '×': 'x',
        'α': 'alpha',
        'β': 'beta',
        'γ': 'gamma',
        'δ': 'delta',
        'λ': 'lambda',
        'μ': 'micro',
        'π': 'pi',
        'σ': 'sigma',
        'τ': 'tau',
        'φ': 'phi',
        'ω': 'omega',
        '≥': '>=',
        '≤': '<=',
        '≠': '!=',
        '∼': '~',
        '′': "'",
        '"': '"',
    }
    for symbol, replacement in replacements.items():
        text = text.replace(symbol, replacement)
    
    # Convert to lowercase and normalize whitespace
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    
    # Normalize unit spacing
    text = re.sub(r'(\d+)\s*(mm|nm|µm|ml|kg|mg|ng|pg|L|ml|μl|°C|Hz|kDa|M|mM|μM|nM|pM|h|min|s|ms|V|mV|A|mA|μA|W|mW|μW)',
                  r'\1 \2', text)
    
    # Normalize scientific notation
    text = re.sub(r'(\d+)[eE]([-+]?\d+)', r'\1 x 10^\2', text)
    
    # Normalize decimal numbers
    text = re.sub(r'(\d+),(\d+)', r'\1.\2', text)
    
    # Normalize ratios
    text = re.sub(r'(\d+)\s*:\s*(\d+)', r'\1:\2', text)
    text = re.sub(r'(\d+)\s*/\s*(\d+)', r'\1/\2', text)
    
    # Normalize percentages
    text = re.sub(r'(\d+)\s*%', r'\1%', text)
    
    # Normalize ranges
    text = re.sub(r'(\d+)\s*-\s*(\d+)', r'\1-\2', text)
    
    # Remove figure and panel labels
    text = re.sub(r'^(figure|fig\.?)\s*\d+[\.:]?\s*', '', text)
    text = re.sub(r'\(?[A-Z]\)?[\.,]?\s+', '', text)
    
    return text.strip()

class FigureCaptionExtractor(ABC):
    """
    Abstract base class for extracting figure captions from a document.

    This class provides common functionality and defines the interface that all
    figure caption extractors should implement.
    """

    def __init__(self, config: Dict):
        """Initialize with configuration."""
        self.config = config

    def _extract_docx_content(self, file_path: str) -> str:
        """
        Extract text content from a DOCX file.

        Args:
            file_path (str): Path to the DOCX file.

        Returns:
            str: Extracted text content from the DOCX file.
        """
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return " ".join(paragraphs)
        except Exception as e:
            logger.exception(f"Error reading DOCX file {file_path}: {str(e)}")
            return ""

    def _parse_response(self, response_text: str) -> Dict[str, str]:
        """
        Parse JSON response containing figure captions.

        Args:
            response_text (str): The response text to parse.

        Returns:
            Dict[str, str]: Dictionary of figure labels and their captions.
        """
        try:
            json_match = re.search(r'```json\s*({[\s\S]*?})\s*```', response_text)
            if json_match:
                json_str = json_match.group(1)
            else:
                json_match = re.search(r'({[^{]*})', response_text)
                if json_match:
                    json_str = json_match.group(1)
                else:
                    json_str = response_text.strip()

            json_str = re.sub(r'[\n\r\t]', ' ', json_str)
            json_str = re.sub(r'\s+', ' ', json_str)

            try:
                captions = json.loads(json_str)
            except json.JSONDecodeError:
                json_str = re.sub(r'([{,])\s*([a-zA-Z0-9_]+)\s*:', r'\1"\2":', json_str)
                captions = json.loads(json_str)

            cleaned_captions = {}
            for key, value in captions.items():
                if isinstance(value, dict) and "title" in value and "caption" in value:
                    title = value["title"].strip()
                    caption = value["caption"].strip()
                    if caption and caption != "Figure caption not found.":
                        clean_key = re.sub(r'^(Fig\.|Figure)\s*', 'Figure ', key)
                        cleaned_captions[clean_key] = {
                            "title": title,
                            "caption": caption
                        }

            return cleaned_captions

        except Exception as e:
            logger.error(f"Error parsing response: {str(e)}", exc_info=True)
            logger.error(f"Error parsing response: {response_text}", exc_info=True)
            import pdb; pdb.set_trace()
            return {}

    def _validate_caption(self, docx_path: str, caption: str, threshold: float = 0.85) -> Tuple[bool, float, str]:
        """
        Validate extracted caption against document text using ROUGE-L score and show differences.
        
        Returns:
            Tuple[bool, float, str]: (validation result, confidence score, diff text)
        """
        try:
            norm_caption = normalize_text(caption)
            if not norm_caption:
                return False, 0.0, ""
            
            doc = Document(docx_path)
            text_blocks = []
            current_block = []
            best_block = ""
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    if current_block:
                        text_blocks.append(" ".join(current_block))
                        current_block = []
                    continue
                    
                if text.lower().startswith(('figure', 'fig.', 'fig ')):
                    if current_block:
                        text_blocks.append(" ".join(current_block))
                    current_block = [text]
                elif current_block:
                    current_block.append(text)
                else:
                    current_block = [text]

            if current_block:
                text_blocks.append(" ".join(current_block))

            best_score = 0.0
            for block in text_blocks:
                norm_block = normalize_text(block)
                if not norm_block:
                    continue

                rouge_scorer_ = rouge_scorer.RougeScorer(['rougeL'], use_stemmer=True)
                scores = rouge_scorer_.score(norm_block, norm_caption)
                rouge_l = scores['rougeL'].fmeasure
                if rouge_l > best_score:
                    best_score = rouge_l
                    best_block = block

            if best_score >= threshold:
                diff = self._generate_diff(best_block, caption)
                return True, best_score, diff
            return False, best_score, ""

        except Exception as e:
            logger.error(f"Error in caption validation: {str(e)}")
            return False, 0.0, ""

    def _generate_diff(self, original: str, ai_generated: str) -> str:
        """Generate a marked-up diff between original and AI-generated text."""
        import difflib
        
        d = difflib.Differ()
        diff = list(d.compare(original.split(), ai_generated.split()))
        
        result = []
        for word in diff:
            if word.startswith('  '):
                result.append(word[2:])
            elif word.startswith('- '):
                result.append(f"\033[91m{word[2:]}\033[0m")  # Red for deletions
            elif word.startswith('+ '):
                result.append(f"\033[92m{word[2:]}\033[0m")  # Green for additions
                
        return ' '.join(result)

    @abstractmethod
    def _locate_figure_captions(self, doc_string: str, expected_figure_count: int, expected_figure_labels: str) -> Optional[str]:
        """
        Locate figure captions in the document using AI.
        
        Args:
            doc_string (str): Document text content
            expected_figure_count (int): Expected number of figures
            expected_figure_labels (str): Expected figure labels
            
        Returns:
            Optional[str]: Located captions text or None
        """
        pass

    @abstractmethod
    def _extract_individual_captions(self, all_captions: str, expected_figure_count: int, expected_figure_labels: str) -> Dict[str, str]:
        """
        Extract individual figure captions from the located captions text.
        
        Args:
            all_captions (str): Text containing all captions
            expected_figure_count (int): Expected number of figures
            expected_figure_labels (str): Expected figure labels
            
        Returns:
            Dict[str, str]: Dictionary of figure labels and their captions
        """
        pass

    @abstractmethod
    def extract_captions(self, docx_path: str, zip_structure: ZipStructure, expected_figure_count: int, expected_figure_labels: str) -> ZipStructure:
        """
        Extract captions from the document and update ZipStructure.
        
        Args:
            docx_path (str): Path to DOCX file
            zip_structure (ZipStructure): Current ZIP structure
            expected_figure_count (int): Expected number of figures
            expected_figure_labels (str): Expected figure labels
            
        Returns:
            ZipStructure: Updated ZIP structure with extracted captions
        """
        pass
    
    @staticmethod
    def normalize_figure_label(label: str) -> str:
        """Normalize figure label to standard format 'Figure X'."""
        # Remove any whitespace and convert to lowercase for comparison
        clean_label = label.strip().lower()
        
        # Extract the figure number
        number = ''.join(filter(str.isdigit, clean_label))
        
        if number:
            return f"Figure {number}"
        return label
